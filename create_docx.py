import docx

doc = docx.Document()
doc.add_heading('Security Implementation Overview', 0)

doc.add_heading('1. Encryption Service', level=1)
doc.add_paragraph('The primary file implementing data encryption in the project is the Encryption Service (lib/services/encryption_service.dart).')

doc.add_heading('What it does:', level=2)
p1 = doc.add_paragraph()
p1.add_run('It provides an ')
p1.add_run('EncryptionService').bold = True
p1.add_run(' class that handles encrypting and decrypting data using AES-256 in CBC mode via the encrypt package.')

doc.add_paragraph('• Key and IV: It currently uses a hardcoded 32-character key and a 16-byte Initialization Vector (IV).', style='List Bullet')
doc.add_paragraph('• Methods: It exposes two static methods: encryptText and decryptText, which appear to be used for securing sensitive data (like chat messages) before sending them to the database.', style='List Bullet')

p2 = doc.add_paragraph()
p2.add_run('Warning:').bold = True
p2.add_run(' As noted in the code, using a hardcoded static key is fine for demonstration purposes, but for a production app, the encryption key should be securely generated, fetched, or derived per user/chat to ensure proper security.')

doc.add_heading('2. Firestore Security Rules', level=1)
doc.add_paragraph('The database security rules and access controls are defined in the firestore.rules file.')

doc.add_heading('Key Rules Implemented:', level=2)

doc.add_heading('Users', level=3)
doc.add_paragraph('• Read: Any authenticated user can read profiles (needed for chats and displaying item owners).', style='List Bullet')
doc.add_paragraph('• Write: Only the user themselves can create or update their own profile. Profile deletion is disabled.', style='List Bullet')

doc.add_heading('Items', level=3)
doc.add_paragraph('• Read/Create: Any authenticated user can view or create items.', style='List Bullet')
doc.add_paragraph('• Update: Only the item\'s owner or the designated receiver can update an item (e.g., to mark it as received/donated).', style='List Bullet')
doc.add_paragraph('• Delete: Only the item\'s owner can delete it.', style='List Bullet')

doc.add_heading('Chats & Messages', level=3)
doc.add_paragraph('• Chats: Users can only create, read, or update a chat if their UID is in the participantIds list.', style='List Bullet')
doc.add_paragraph('• Messages: A user can only read or send messages in a chat if they are verified as a participant in that specific chat room. Messages cannot be edited or deleted once sent.', style='List Bullet')

doc.add_heading('Admin Actions', level=3)
doc.add_paragraph('• Creation: Users can create feedback and reports (tied to their UID).', style='List Bullet')
doc.add_paragraph('• Access Restrictions: Only users with the \'Admin\' role in their user document can read/update these administrative documents, except for the users themselves who can read their own feedback/reports.', style='List Bullet')

doc.add_heading('3. Authentication (Other Security Files)', level=1)
doc.add_paragraph('The project also implements user authentication and identity security in the lib/screens/auth/ directory. Key files include:')
doc.add_paragraph('• login_screen.dart', style='List Bullet')
doc.add_paragraph('• signup_screen.dart', style='List Bullet')
doc.add_paragraph('• phone_auth_screen.dart', style='List Bullet')
doc.add_paragraph('• otp_screen.dart', style='List Bullet')

doc.save('Security_Implementation.docx')
print("Successfully created Security_Implementation.docx")
